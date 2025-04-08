import os
import json
import asyncio
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import List, Dict, Any
from dotenv import load_dotenv
from supabase import create_client, Client

# Import for PDF and DOCX processing
import PyPDF2
from docx import Document

# Import SentenceTransformer and numpy for embeddings
from sentence_transformers import SentenceTransformer
import numpy as np

# Load environment variables
load_dotenv()

# Initialize Supabase client
supabase: Client = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_SERVICE_KEY")
)

# Initialize SentenceTransformer model (all-mpnet-base-v2 returns 768-dim embeddings)
embedding_model = SentenceTransformer('all-mpnet-base-v2')

#########################################
# Text Extraction Functions
#########################################
def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from a resume file.
    Supports PDF (via PyPDF2), DOCX (via python-docx), and TXT.
    """
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith('.txt'):
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF using PyPDF2.
    """
    text = ""
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text.strip()

#########################################
# Utility Functions
#########################################
def extract_summary(text: str) -> str:
    """
    Extract a simple summary from the text.
    This implementation returns the first sentence (or a short snippet).
    """
    sentences = text.split('.')
    if sentences and sentences[0].strip():
        return sentences[0].strip() + '.'
    else:
        return text[:150] + '...' if len(text) > 150 else text

def chunk_text(text: str, chunk_size: int = 5000) -> List[str]:
    """
    Splits large text into smaller chunks.
    Attempts to break at paragraph boundaries.
    """
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size

        if end >= text_length:
            chunks.append(text[start:].strip())
            break

        # Try to break at a double newline for paragraph boundaries
        chunk = text[start:end]
        last_break = chunk.rfind('\n\n')
        if last_break != -1 and last_break > chunk_size * 0.3:
            end = start + last_break

        chunks.append(text[start:end].strip())
        start = end

    return chunks

async def get_embedding(text: str) -> List[float]:
    """
    Uses SentenceTransformers to generate an embedding vector for the provided text.
    This function wraps the synchronous model.encode call in an executor.
    """
    loop = asyncio.get_event_loop()
    # Run the encoding in a separate thread to avoid blocking
    embedding = await loop.run_in_executor(None, embedding_model.encode, text)
    if isinstance(embedding, np.ndarray):
        return embedding.tolist()
    else:
        return list(embedding)

#########################################
# Data structure for a processed resume chunk
#########################################
@dataclass
class ProcessedResumeChunk:
    url: str              # Using the file name or URL
    chunk_number: int
    title: str            # Derived from the file name (without extension)
    summary: str          # A simple summary of the chunk
    content: str          # The actual text content of this chunk
    metadata: Dict[str, Any]
    embedding: List[float]

#########################################
# Processing and Insertion Functions
#########################################
async def process_resume_chunk(chunk: str, chunk_number: int, file_name: str) -> ProcessedResumeChunk:
    """
    Process a single chunk of a resume:
      - Extract a summary using a heuristic.
      - Generate an embedding using SentenceTransformers.
      - Prepare metadata and other fields according to the schema.
    """
    summary = extract_summary(chunk)
    embedding = await get_embedding(chunk)
    metadata = {
        "file_name": file_name,
        "chunk_size": len(chunk),
        "processed_at": datetime.now(timezone.utc).isoformat()
    }
    # Use the file name (without extension) as the title.
    title = os.path.splitext(file_name)[0]
    return ProcessedResumeChunk(
        url=file_name,         # Using file_name as the URL; adjust as needed.
        chunk_number=chunk_number,
        title=title,
        summary=summary,
        content=chunk,         # The column "content" matches your schema.
        metadata=metadata,
        embedding=embedding
    )

def insert_resume_chunk(chunk: ProcessedResumeChunk):
    """
    Insert the processed resume chunk into the Supabase 'resumes' table.
    """
    try:
        data = {
            "url": chunk.url,
            "chunk_number": chunk.chunk_number,
            "title": chunk.title,
            "summary": chunk.summary,
            "content": chunk.content,
            "metadata": chunk.metadata,
            "embedding": chunk.embedding
        }
        result = supabase.table("resumes").insert(data).execute()
        print(f"Inserted chunk {chunk.chunk_number} for resume {chunk.url}")
        return result
    except Exception as e:
        print(f"Error inserting chunk {chunk.chunk_number} for resume {chunk.url}: {e}")
        return None

async def process_resume_file(file_path: str):
    """
    Process a resume file:
      1. Extract text from the file.
      2. Chunk the text if the resume is large.
      3. Process each chunk concurrently.
      4. Insert each processed chunk into the database.
    """
    file_name = os.path.basename(file_path)
    resume_text = extract_text_from_resume(file_path)
    print(f"Processing resume file: {file_name}")

    chunks = chunk_text(resume_text, chunk_size=5000)
    print(f"Split resume into {len(chunks)} chunks")

    tasks = [
        process_resume_chunk(chunk, idx, file_name)
        for idx, chunk in enumerate(chunks)
    ]
    processed_chunks = await asyncio.gather(*tasks)

    for chunk in processed_chunks:
        insert_resume_chunk(chunk)

async def process_all_resumes(directory: str):
    """
    Loops through all resume files (.pdf, .docx, or .txt) in a directory,
    processes them, and stores the results.
    """
    resume_files = [
        os.path.join(directory, f)
        for f in os.listdir(directory)
        if f.lower().endswith(('.pdf', '.docx', '.txt'))
    ]
    tasks = [process_resume_file(file_path) for file_path in resume_files]
    await asyncio.gather(*tasks)

#########################################
# Main Entry Point
#########################################
if __name__ == "__main__":
    resumes_directory = "P:\\AI\\ottomator-agents\\crawl4AI-agent\\resumes"
    asyncio.run(process_all_resumes(resumes_directory))import os
import json
import asyncio
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import List, Dict, Any
from dotenv import load_dotenv
from supabase import create_client, Client

# Import for PDF and DOCX processing
import PyPDF2
from docx import Document

# Import SentenceTransformer and numpy for embeddings
from sentence_transformers import SentenceTransformer
import numpy as np

# Load environment variables
load_dotenv()

# Initialize Supabase client
supabase: Client = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_SERVICE_KEY")
)

# Initialize SentenceTransformer model (all-mpnet-base-v2 returns 768-dim embeddings)
embedding_model = SentenceTransformer('all-mpnet-base-v2')


#########################################
# Text Extraction Functions
#########################################
def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from a resume file.
    Supports PDF (via PyPDF2), DOCX (via python-docx), and TXT.
    """
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith('.txt'):
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF using PyPDF2.
    """
    text = ""
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text.strip()


#########################################
# Utility Functions
#########################################
def extract_summary(text: str) -> str:
    """
    Extract a simple summary from the text.
    This implementation returns the first sentence (or a short snippet).
    """
    sentences = text.split('.')
    if sentences and sentences[0].strip():
        return sentences[0].strip() + '.'
    else:
        return text[:150] + '...' if len(text) > 150 else text

def chunk_text(text: str, chunk_size: int = 5000) -> List[str]:
    """
    Splits large text into smaller chunks.
    Attempts to break at paragraph boundaries.
    """
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size

        if end >= text_length:
            chunks.append(text[start:].strip())
            break

        # Try to break at a double newline for paragraph boundaries
        chunk = text[start:end]
        last_break = chunk.rfind('\n\n')
        if last_break != -1 and last_break > chunk_size * 0.3:
            end = start + last_break

        chunks.append(text[start:end].strip())
        start = end

    return chunks

async def get_embedding(text: str) -> List[float]:
    """
    Uses SentenceTransformers to generate an embedding vector for the provided text.
    This function wraps the synchronous model.encode call in an executor.
    """
    loop = asyncio.get_event_loop()
    # Run the encoding in a separate thread to avoid blocking
    embedding = await loop.run_in_executor(None, embedding_model.encode, text)
    if isinstance(embedding, np.ndarray):
        return embedding.tolist()
    else:
        return list(embedding)


#########################################
# Data structure for a processed resume chunk
#########################################
@dataclass
class ProcessedResumeChunk:
    url: str              # Using the file name or URL
    chunk_number: int
    title: str            # Derived from the file name (without extension)
    summary: str          # A simple summary of the chunk
    content: str          # The actual text content of this chunk
    metadata: Dict[str, Any]
    embedding: List[float]


#########################################
# Processing and Insertion Functions
#########################################
async def process_resume_chunk(chunk: str, chunk_number: int, file_name: str) -> ProcessedResumeChunk:
    """
    Process a single chunk of a resume:
      - Extract a summary using a heuristic.
      - Generate an embedding using SentenceTransformers.
      - Prepare metadata and other fields according to the schema.
    """
    summary = extract_summary(chunk)
    embedding = await get_embedding(chunk)
    metadata = {
        "file_name": file_name,
        "chunk_size": len(chunk),
        "processed_at": datetime.now(timezone.utc).isoformat()
    }
    # Use the file name (without extension) as the title.
    title = os.path.splitext(file_name)[0]
    return ProcessedResumeChunk(
        url=file_name,         # Using file_name as the URL; adjust as needed.
        chunk_number=chunk_number,
        title=title,
        summary=summary,
        content=chunk,         # The column "content" matches your schema.
        metadata=metadata,
        embedding=embedding
    )

def insert_resume_chunk(chunk: ProcessedResumeChunk):
    """
    Insert the processed resume chunk into the Supabase 'resumes' table.
    """
    try:
        data = {
            "url": chunk.url,
            "chunk_number": chunk.chunk_number,
            "title": chunk.title,
            "summary": chunk.summary,
            "content": chunk.content,
            "metadata": chunk.metadata,
            "embedding": chunk.embedding
        }
        result = supabase.table("resumes").insert(data).execute()
        print(f"Inserted chunk {chunk.chunk_number} for resume {chunk.url}")
        return result
    except Exception as e:
        print(f"Error inserting chunk {chunk.chunk_number} for resume {chunk.url}: {e}")
        return None

async def process_resume_file(file_path: str):
    """
    Process a resume file:
      1. Extract text from the file.
      2. Chunk the text if the resume is large.
      3. Process each chunk concurrently.
      4. Insert each processed chunk into the database.
    """
    file_name = os.path.basename(file_path)
    resume_text = extract_text_from_resume(file_path)
    print(f"Processing resume file: {file_name}")

    chunks = chunk_text(resume_text, chunk_size=5000)
    print(f"Split resume into {len(chunks)} chunks")

    tasks = [
        process_resume_chunk(chunk, idx, file_name)
        for idx, chunk in enumerate(chunks)
    ]
    processed_chunks = await asyncio.gather(*tasks)

    for chunk in processed_chunks:
        insert_resume_chunk(chunk)

async def process_all_resumes(directory: str):
    """
    Loops through all resume files (.pdf, .docx, or .txt) in a directory,
    processes them, and stores the results.
    """
    resume_files = [
        os.path.join(directory, f)
        for f in os.listdir(directory)
        if f.lower().endswith(('.pdf', '.docx', '.txt'))
    ]
    tasks = [process_resume_file(file_path) for file_path in resume_files]
    await asyncio.gather(*tasks)

#########################################
# Main Entry Point
#########################################
if __name__ == "__main__":
    resumes_directory = "P:\\AI\\ottomator-agents\\crawl4AI-agent\\resumes"
    asyncio.run(process_all_resumes(resumes_directory))
